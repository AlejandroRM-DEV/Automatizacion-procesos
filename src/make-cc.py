import docx
import re
import openpyxl
import argparse
import os


def get_no_response(doc):
    header = doc.sections[0].header
    header_text = ''
    for paragraph in header.paragraphs:
        header_text += paragraph.text

    match = re.search(
        r'D-(X{0,3}(I{0,3}|IV|IX|V?I{0,3})|[IVX]{1,2})/\d{1,6}/\d{4}/IJCF/0*\d{6}/\d{4}/IF/0*\d{2}', header_text)

    return match.group() if match else ''


def get_evidence_table(doc):
    content = ""
    for table in doc.tables:
        if table.rows[0].cells[0].text == "INDICIO":
            for row in table.rows[1:]:
                for cell in row.cells:
                    content += cell.text + " "
                content += "\r\n"

    return content.rstrip()


def get_prosecutor(doc):
    return doc.paragraphs[0].text


def get_agency(doc):
    agency = ""
    for paragraph in doc.paragraphs[2:]:
        if "P R E S E N T E" in paragraph.text:
            break
        else:
            agency += paragraph.text + " "

    return agency


def post_process(data):
    parts = data["no_response"].split("/")

    data["ci"] = parts[1] + "/" + parts[2]

    if parts[7] == "01":
        data["type"] = "ADQUISICIÓN DE DATOS DE TELÉFONO CELULAR (DISPOSITIVOS MOVILES)"
    elif parts[7] == "05":
        data["type"] = "INVESTIGACIÓN SOBRE SERVICIOS Y APLICACIONES DE INTERNET"
    elif parts[7] == "07":
        data["type"] = "EXTRACCIÓN DE INFORMACIÓN DE DISPOSITIVOS DE ALMACENAMIENTO DIGITAL"
    elif parts[7] == "10":
        data["type"] = "EXTRACCIÓN DE INFORMACIÓN DE EQUIPOS DE GRABACIÓN DE VIDEO"
    elif parts[7] == "11":
        data["type"] = "ADQUISICIÓN DE INFORMACIÓN DE EQUIPOS DE CÓMPUTO"
    elif parts[7] == "12":
        data["type"] = "IDENTIFICACIÓN TÉCNICA DE VIDEOS Y SECUENCIA"
    else:
        data["type"] = ""

    return data


def write_chain_of_custody(dest_directory):
    script_dir = os.path.dirname(os.path.abspath(__file__))
    path_cc = os.path.abspath(f"{script_dir}/protected/templates/cc.xlsx")
    workbook = openpyxl.load_workbook(path_cc)

    hoja = workbook['FORMATO']
    hoja['H45'] = data["evidence_table"]
    hoja['F45'] = data["evidence_table"].count("Número de serie de volumen")
    hoja['K9'] = data["agency"]
    hoja['V9'] = data["prosecutor"]
    hoja['AA6'] = data["ci"]
    hoja['F12'] = data["type"]

    workbook.save(os.path.abspath(f"{dest_directory}/cc.xlsx"))
    os.startfile(os.path.abspath(f"{dest_directory}/cc.xlsx"))


def write_envelope(dest_directory):
    script_dir = os.path.dirname(os.path.abspath(__file__))
    path_cc = os.path.abspath(f"{script_dir}/protected/templates/sobre.xlsx")
    workbook = openpyxl.load_workbook(path_cc)

    hoja = workbook['Hoja1']
    hoja['D16'] = data["no_response"]
    hoja['D18'] = data["ci"]
    hoja['D22'] = data["agency"]
    hoja['C26'] = data["evidence_table"]

    workbook.save(os.path.abspath(f"{dest_directory}/sobre.xlsx"))
    os.startfile(os.path.abspath(f"{dest_directory}/sobre.xlsx"))


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("-f", "--file", help="Archivo", required=True)
    args = parser.parse_args()

    doc = docx.Document(args.file)
    data = {}
    data["no_response"] = get_no_response(doc)
    data["evidence_table"] = get_evidence_table(doc)
    data["agency"] = get_agency(doc)
    data["prosecutor"] = get_prosecutor(doc)
    data = post_process(data)

    dest_directory = os.path.dirname(args.file)
    write_chain_of_custody(dest_directory)
    write_envelope(dest_directory)
